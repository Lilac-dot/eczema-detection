"""Generate a comprehensive Word report covering the model and the full project status."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\tishy\Documents\Honors\AD_Wearable_Model_and_Project_Report_2026-08-26.docx"
DOCS = r"C:\Users\tishy\Documents\Honors\docs"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table_from_rows(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def picture(filename, width=4.5, caption=None):
    doc.add_picture(f"{DOCS}\\{filename}", width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


# ===================== TITLE =====================
doc.add_heading("AD Wearable Project — Model & Progress Report", level=0)
p("Report date: 26 August 2026", italic=True)
p("Multi-Sensor Wearable System for Monitoring Eczema Severity", italic=True)

doc.add_paragraph()
h1("1. Executive Summary")
p(
    "This report covers the model-building and dataset-validation work carried out across "
    "this project session, extending the earlier Progress Report and Progress Log. Two "
    "major threads were pursued: Stage A (motion-based scratch detection from wearable "
    "sensors) and Stage B (image-based eczema diagnosis). Stage B saw the most significant "
    "progress: a serious shortcut-learning problem was discovered, diagnosed, and fixed "
    "through a multi-step investigation, ending in a genuinely trustworthy model. Stage A "
    "was found to face a hard sensor-hardware limitation that no amount of data-sourcing "
    "or modelling can bypass. Stage C (fusion) remains not yet started, pending decisions "
    "on Stage A's direction."
)
table_from_rows(
    ["Stage", "Status", "Headline result"],
    [
        ["A — Motion/Scratch", "First proxy model built; hardware ceiling identified", "Test AUC 0.83 (WISDM proxy, teeth-brushing vs. other)"],
        ["B — Image Diagnosis", "Shortcut found, diagnosed, and fixed", "81.07% accuracy / F1 81.18% (Eczema vs. similar diseases, balanced)"],
        ["C — Fusion", "Not started", "Blocked on Stage A direction"],
    ],
)

# ===================== SECTION 2: STAGE B =====================
h1("2. Stage B — Image-Based Eczema Diagnosis")

h2("2.1 Starting point and first models")
p(
    "The original dataset (dataset/Eczema/, dataset/Normal/, ~2,757 images after cleaning "
    "104 exact and 262 near-duplicates) produced two first-pass models: a ResNet18 CNN "
    "(95.66% test accuracy) and a colour-features + LightGBM model (95.18% test accuracy), "
    "following Maulana et al. (2024)'s approach."
)

h2("2.2 Diagnosis: shortcut learning")
p(
    "Both models scoring within half a point of each other at a suspiciously high 95% was "
    "the first red flag. Direct inspection showed Eczema images were clinical DermNet "
    "macro photos (many watermarked) while Normal images were largely unrelated stock "
    "photography. Quantified: mean whole-image brightness (MEAN_XYZ_2) was 0.15 for Eczema "
    "vs. 0.58 for Normal — a ~4x gap large enough on its own to explain the accuracy, "
    "confirmed as LightGBM's single most important feature by a wide margin."
)

h2("2.3 Two candidate fixes tested and ruled out")
bullet("Brightness/exposure normalization (forcing identical per-image mean/std): accuracy only dropped from 95.18% to 93.98% — barely moved.")
bullet("Auto-cropping to the skin/lesion region only: accuracy did not change at all (CNN 95.66%→95.66% identical; LightGBM 95.18%→95.42%). The brightness gap survived cropping essentially unchanged (0.154 vs. 0.494).")
p(
    "Conclusion: the confound is baked into the entire image-capture pipeline (camera, "
    "lighting, compression) for each source, not just composition or exposure — neither is "
    "fixable by preprocessing alone.", bold=True
)
picture("confusion_matrix.png", caption="Original Eczema-vs-Normal CNN — the 95.66% figure now known to be shortcut-inflated.")

h2("2.4 The real fix: reframing as Eczema vs. similar diseases")
p(
    "A second dataset (SkinDisease/, Kaggle 'Human Skin Diseases', 20 classes) was found "
    "and cleaned (corruption/duplicate check across 17,266 images; also checked "
    "specifically for cross-split duplicates, finding 173 test-vs-train/valid leaks and "
    "two classes whose entire 'valid' folder was 100% duplicate of their 'train' folder). "
    "Its 'Normal' class was found badly contaminated (unrelated stock/product photos — a "
    "drill kit, cherries, a cat — quantified via a skin-colour-fraction screen showing "
    "12-16.6% of its images had essentially no skin-coloured pixels, 2-8x worse than any "
    "disease class) and internally inconsistent across its own splits. It was excluded "
    "entirely."
)
p(
    "The disease classes, however, were confirmed same-source (DermNet-style, matching "
    "the original dataset's own Eczema images almost exactly by filename convention). A "
    "curated 7-class 'other disease' set was selected — Psoriasis, Tinea, Candidiasis, "
    "Infestations_Bites, Lichen, DrugEruption, Rosacea — chosen specifically because they "
    "are genuine clinical look-alikes for eczema, unlike unrelated classes (Vitiligo, "
    "Moles, Skin Cancer) that would make the task trivially easy again."
)

h2("2.5 Iteration history")
table_from_rows(
    ["Version", "Change", "Test Accuracy", "F1 (Eczema)"],
    [
        ["v1 CNN", "Frozen backbone (only final layer trained)", "69.55%", "46.74%"],
        ["v1 LightGBM", "Colour features only", "82.43%", "57.06%"],
        ["v2 CNN", "Fine-tuned layer4 + fc, softened class weight", "87.00%", "63.01%"],
        ["Ensemble (v2)", "Average of CNN v2 + LightGBM", "89.05%", "67.84%"],
        ["Balanced CNN (final)", "Merged 3rd Eczema source, rebalanced ~50/50", "81.07%", "81.18%"],
    ],
)
p(
    "The v1 CNN underperforming LightGBM (opposite of the original shortcut result, where "
    "both models agreed) was itself a good sign — it meant there was no shared shortcut "
    "left for both models to find. Fine-tuning more of the network (v2) then closed that "
    "gap and overtook LightGBM, as expected once given enough capacity."
)

h2("2.6 Final step: merging a third source and rebalancing")
p(
    "A third dataset (Eczema/, 17 DermNet-subtype folders, 1,395 images) was added and "
    "checked for genuinely new content. Only 5 of 719 cleaned, on-topic images were "
    "actually new — 714 were duplicates of what the project already had, confirming this "
    "is largely the same underlying archive as the other two sources. Curating which "
    "subfolders counted as Eczema mattered: Ichthyosis, Keratosis pilaris, Neurotic "
    "excoriations, Prurigo nodularis, Keratolysis exfoliativa, and Lichen simplex "
    "chronicus (which would have directly contradicted the existing 'Lichen' "
    "other-disease class) were excluded as distinct diagnoses."
)
p(
    "Final unique Eczema pool: 1,665 images. Since there was not enough real Eczema data "
    "to reach 50/50 by growing that side, the 'Other' pool was downsampled from 4,460 to "
    "~1,665 instead (proportionally across all 7 diseases) — an honest trade of total "
    "data volume (5,494 → 3,330 images) for genuine class balance."
)

h2("2.7 Final result")
table_from_rows(
    ["Metric", "Value"],
    [
        ["Test set size", "507 images"],
        ["Test accuracy", "81.07%"],
        ["Precision (Eczema)", "79.92%"],
        ["Recall (Eczema)", "82.47%"],
        ["F1 (Eczema)", "81.18%"],
    ],
)
picture("confusion_matrix_balanced.png", caption="Final Stage B model (balanced CNN) — test set confusion matrix.")
p(
    "Raw accuracy is lower than the imbalanced-dataset versions (89.05%), but this is the "
    "correct direction, not a regression: the earlier higher accuracy was partly propped "
    "up by class imbalance. With genuine 50/50 balance, accuracy and F1 now sit close "
    "together (81.07% vs. 81.18%), and precision/recall are balanced (80%/82%) rather "
    "than lopsided — a sign this number is not being inflated. The model's errors "
    "consistently cluster on Psoriasis, Infestations_Bites, Tinea, and Lichen across "
    "every version tested — the genuine real-world clinical look-alikes for eczema — "
    "which is evidence the model is tracking real visual similarity, not noise.", bold=True
)

# ===================== SECTION 3: HOW THE MODEL WORKS =====================
h1("3. How the Model Works")

h2("3.1 Labels")
p(
    "Each image carries one binary label: 1 = Eczema, 0 = Other (a pooled bucket of the "
    "7 curated look-alike diseases). Each image is represented as a 224x224x3 grid of "
    "pixel values (roughly 150,000 numbers) for the CNN, or as 90 hand-computed colour "
    "statistics (means/standard deviations across 15 colour spaces) for the LightGBM model."
)

h2("3.2 The CNN (ResNet18)")
p(
    "A chain of convolutional filters scans the image for increasingly complex patterns "
    "(edges -> textures -> lesion-like structures), ending in two raw scores (one per "
    "class). These are converted to probabilities via softmax:"
)
p("P(class i) = exp(z_i) / (exp(z_Eczema) + exp(z_Other))", italic=True)
p(
    "Training minimises cross-entropy loss, loss = -log(P(correct class)) — a confident, "
    "correct prediction produces a near-zero penalty, while a confident, wrong prediction "
    "is penalised heavily. Backpropagation computes, for every internal weight, how much "
    "nudging it up or down would reduce this loss, and gradient descent applies a small "
    "step in that direction, repeated over many images and passes (epochs)."
)
p(
    "Transfer learning: rather than training from scratch, the network started from "
    "weights pretrained on millions of general (non-medical) photos. Most of the network "
    "was kept fixed ('frozen'); only the final residual block (layer4) and the final "
    "decision layer were allowed to adapt, using a 10x smaller learning rate on layer4 "
    "than on the final layer — enough adaptation to learn task-specific texture cues "
    "without erasing the useful general vision features already learned."
)

h2("3.3 Evaluation metrics")
p("Using TP/FP/FN (true positives, false positives, false negatives):")
p("Precision = TP / (TP + FP)      Recall = TP / (TP + FN)      F1 = 2 x Precision x Recall / (Precision + Recall)", italic=True)
p(
    "F1 is the harmonic mean of precision and recall — it cannot be gamed by maximising "
    "one at the expense of the other, which is why it was tracked throughout rather than "
    "accuracy alone, especially once class imbalance was a factor."
)

# ===================== SECTION 4: STAGE A =====================
h1("4. Stage A — Motion-Based Scratch Detection")

h2("4.1 First model")
p(
    "WISDM (51 subjects, phone/watch accelerometer + gyroscope, cleaned to 15.36M rows) "
    "was used to build a first proxy model: a LightGBM classifier on WISDM's own "
    "pre-extracted phone-accelerometer features (91 features/window), predicting "
    "'teeth-brushing' (the closest available WISDM proxy for repetitive hand motion) vs. "
    "everything else. Subject-level train/val/test split (not row-level, to prevent "
    "leakage) with AUC-based early stopping and a validation-tuned decision threshold."
)
table_from_rows(
    ["Metric", "Validation", "Test"],
    [
        ["AUC", "0.878", "0.827"],
        ["Accuracy (tuned threshold)", "89.96%", "92.33%"],
        ["F1 (teeth)", "37.54%", "25.07%"],
    ],
)
picture("confusion_matrix_stage_a.png", caption="Stage A first model — subject-held-out test confusion matrix.")

h2("4.2 A literature-grounded hardware limitation")
p(
    "A paper the user provided (Chun et al. 2021, Science Advances — the 'ADAM' sensor) "
    "showed that the signal which actually distinguishes scratching from other hand "
    "motion is a 100-800 Hz acousto-mechanic vibration, and that even a 100 Hz-sampling "
    "smartwatch cannot reliably capture it (their benchmark smartwatch algorithm confused "
    "hand-waving with scratching for exactly this reason). WISDM samples at only 20 Hz — "
    "5x below that inadequate smartwatch, 80x below the ADAM sensor's 1600 Hz. This "
    "reframes Stage A's modest performance: it is not simply a wrong-proxy-label problem, "
    "but a hardware bandwidth ceiling no amount of relabelling or modelling can fix.", bold=True
)

h2("4.3 Search for a better dataset")
p(
    "A systematic review of 21 scratch-detection studies (checked directly) confirmed "
    "none has released a public dataset — all say data is available on request at best, "
    "citing participant privacy. A general public sleep-state dataset (Child Mind "
    "Institute, Kaggle) was also checked as a possible reframing (using sleep disruption "
    "as an itch proxy instead of scratch motion directly) but was found not to resolve the "
    "core problem either: it has no itch or AD-patient labels at all, so it cannot "
    "validate whether any detected restlessness is actually caused by itch versus any of "
    "the many other causes of disrupted sleep."
)
p(
    "Conclusion: every public-data option checked hits the same wall. A trustworthy Stage "
    "A result requires real patient data pairing wearable motion with genuine itch/AD "
    "severity ground truth — which requires the IRB step already on this project's "
    "critical path.", bold=True
)

# ===================== SECTION 5: STAGE C =====================
h1("5. Stage C — Fusion Model")
p(
    "Not yet started. Design is blocked on Stage A's direction being settled (proxy "
    "pipeline demonstration vs. pursuing real patient/scratch data) and on Stage B's "
    "output format being finalised (a severity score vs. the current binary "
    "Eczema-vs-other-disease framing)."
)

# ===================== SECTION 6: STATUS AND NEXT STEPS =====================
h1("6. Current Status and Recommended Next Steps")
h2("6.1 Decisions needed")
bullet("Stage A direction: ship the WISDM proxy as a documented pipeline demonstration, request data from SIGMA/ADAM authors, and/or scope a self-collected pilot (needs IRB either way).")
bullet("Start the IRB conversation now, in parallel with everything else — it is the slowest-moving item and gates both real patient data and any self-collected pilot.")

h2("6.2 Not yet done")
bullet("Stage C fusion model design and implementation.")
bullet("Sanity-check the 7-class 'similar disease' list with someone with real dermatology expertise.")
bullet("Writing pass integrating the shortcut-learning diagnosis and Stage A's sampling-rate limitation as explicit methodological content, not just caveats.")

doc.save(OUT)
print(f"Saved: {OUT}")
