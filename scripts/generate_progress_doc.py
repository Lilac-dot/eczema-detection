"""Generate a Word document summarizing the work-to-date on the AD wearable project."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\tishy\Documents\Honors\Progress_Log_2026-08-26.docx"

doc = Document()

# --- base style ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def table_from_rows(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# ===================== TITLE =====================
title = doc.add_heading("AD Wearable Project — Progress Log", level=0)
p("Session date: 26 August 2026", italic=True)
p("Prepared for: Lilac (honors project — Multi-Sensor Wearable System for Monitoring Eczema Severity)", italic=True)

doc.add_paragraph()
h1("1. Purpose of This Log")
p(
    "This document records the practical data-preparation and modelling work carried out in this "
    "session, following on from the project's Progress Report and Literature Review "
    "(AD_Wearable_Project_Report.docx). That report set two immediate next steps: (1) download and "
    "confirm public datasets load/inspect correctly, and (2) build a minimal end-to-end pipeline "
    "skeleton. This session addressed both for Stage B (image severity model) and made a start on "
    "Stage A (motion/scratch model) by acquiring and cleaning a motion dataset."
)

# ===================== SECTION 2: STAGE B IMAGE DATASET =====================
h1("2. Stage B — Eczema Image Dataset")

h2("2.1 Starting state")
p(
    "An image dataset was already present at dataset/Eczema/ (1,544 files) and dataset/Normal/ "
    "(1,579 files), matching the 'Eczema Infected + Normal' Kaggle dataset referenced in Section 4.2 "
    "of the project report (3,000+ labelled eczema vs. normal skin images)."
)

h2("2.2 Cleaning process")
p("A cleaning script (scripts/clean_dataset.py) was written and run over both folders. For every image it:")
bullet("Verified the file actually opens as a valid image (Pillow) — catches corrupt/truncated files.")
bullet("Flagged images below 64x64 px as likely icons/thumbnails rather than real photos.")
bullet("Detected exact duplicates via MD5 hash.")
bullet("Detected near-duplicates via perceptual hash (hamming distance <= 4) — catches renamed copies of the same clinical photo.")
p("Nothing was deleted. Flagged files were moved to dataset/_quarantine/<class>/<reason>/ so the decision is fully reversible.")

h2("2.3 Results")
table_from_rows(
    ["Metric", "Value"],
    [
        ["Images scanned", "3,123"],
        ["Corrupt / unreadable", "0"],
        ["Too small (<64px)", "0"],
        ["Exact duplicates", "104"],
        ["Near-duplicates", "262"],
        ["Total quarantined", "366"],
        ["Clean set — Eczema", "1,428"],
        ["Clean set — Normal", "1,329"],
        ["Clean set — Total", "2,757"],
    ],
)
p(
    "Outputs: dataset/manifest_clean.csv (surviving files + labels), dataset/clean_report.csv "
    "(every flagged file + reason)."
)

h2("2.4 Train/val/test split")
p(
    "A stratified 70/15/15 split (scripts/split_dataset.py, seed=42) was applied to the clean "
    "manifest, producing dataset/manifest_train.csv, manifest_val.csv, manifest_test.csv:"
)
table_from_rows(
    ["Split", "Eczema", "Normal", "Total"],
    [
        ["Train", 999, 930, 1929],
        ["Validation", 214, 199, 413],
        ["Test", 215, 200, 415],
    ],
)

# ===================== SECTION 3: STAGE B MODEL =====================
h1("3. Stage B — Image Classifier (First Model)")

h2("3.1 Approach")
p(
    "Transfer learning was selected over training from scratch, given the dataset size (~2.7k "
    "images). A ResNet18 backbone pretrained on ImageNet was used, with the convolutional layers "
    "frozen and only the final fully-connected layer retrained for the binary Eczema-vs-Normal task. "
    "This is a standard, fast, and reproducible baseline appropriate for a first iteration."
)
p("Training configuration:")
bullet("Backbone: ResNet18 (ImageNet1K_V1 weights), backbone frozen, FC layer fine-tuned")
bullet("Input size: 224x224, ImageNet normalization")
bullet("Augmentation (train only): random horizontal flip, random rotation (15 degrees), color jitter")
bullet("Optimizer: Adam, lr=1e-4, on FC parameters only")
bullet("Loss: cross-entropy; Batch size: 32; Epochs: 10; Device: CPU")
p("Scripts: scripts/train_stage_b.py (training), scripts/eval_stage_b.py (test-set evaluation).")

h2("3.2 Training results (validation, per epoch)")
table_from_rows(
    ["Epoch", "Train Loss", "Train Acc", "Val Loss", "Val Acc"],
    [
        [1, 0.6052, "67.86%", 0.4716, "81.11%"],
        [2, 0.4111, "89.11%", 0.3474, "91.04%"],
        [3, 0.3117, "93.62%", 0.2737, "94.19%"],
        [4, 0.2495, "94.82%", 0.2178, "95.64%"],
        [5, 0.2124, "95.90%", 0.1942, "95.88%"],
        [6, 0.1876, "96.16%", 0.1676, "96.13%"],
        [7, 0.1656, "96.37%", 0.1442, "97.34%"],
        [8, 0.1500, "96.94%", 0.1352, "96.85%"],
        [9, 0.1427, "96.63%", 0.1203, "97.58%"],
        [10, 0.1323, "96.63%", 0.1118, "97.58%"],
    ],
)
p("Best validation accuracy: 97.58% (epoch 9). Model checkpoint saved to models/stage_b_resnet18.pt.", bold=True)

h2("3.3 Held-out test set results")
table_from_rows(
    ["Metric", "Value"],
    [
        ["Test set size", "415 images"],
        ["Test accuracy", "95.66%"],
        ["Precision (Eczema)", "92.64%"],
        ["Recall (Eczema)", "99.53%"],
        ["F1 (Eczema)", "95.96%"],
    ],
)
p("Confusion matrix (rows = true label, columns = predicted label):")
table_from_rows(
    ["", "Predicted Normal", "Predicted Eczema"],
    [
        ["True Normal", "183 (TN)", "17 (FP)"],
        ["True Eczema", "1 (FN)", "214 (TP)"],
    ],
)
doc.add_picture(r"C:\Users\tishy\Documents\Honors\docs\confusion_matrix.png", width=Inches(4.0))
p(
    "Interpretation: the model rarely misses eczema (only 1 false negative out of 215 eczema test "
    "images), at the cost of a modest false-positive rate on Normal images (17/200). For a "
    "screening/monitoring use case, that trade-off (high recall, slightly lower precision) is the "
    "safer direction to err in. However, see Section 3.4 below — this accuracy figure is very "
    "likely inflated by a dataset shortcut rather than genuine eczema detection, and should not be "
    "reported as-is."
)

h2("3.4 Data validity concern — likely shortcut learning (important)")
para = p("This accuracy figure should be treated with caution before it is reported anywhere.", bold=True)
p(
    "On inspecting sample images from each class directly, a systematic difference in photo source "
    "and composition was found between the two classes, independent of skin condition:"
)
bullet("Normal/ images are largely stock photography (filenames follow the Shutterstock preview "
       "naming pattern, e.g. '...-260nw-2335137965.jpg'): smiling models, full face/body framing, "
       "colored studio or outdoor backgrounds.")
bullet("Eczema/ images are clinical macro close-ups of lesions on plain/dark backgrounds; a "
       "meaningful fraction (150 of 1,428 files, ~10.5% by filename alone, likely more by visual "
       "style) carry a visible '\u00a9DermNet.com' watermark baked into the image.")
p("Representative examples pulled directly from the dataset:")

# side-by-side example images via a 1x2 table
ex_table = doc.add_table(rows=2, cols=2)
ex_table.autofit = True
c0, c1 = ex_table.rows[0].cells
c0.paragraphs[0].add_run().add_picture(r"C:\Users\tishy\Documents\Honors\docs\example_normal.jpg", width=Inches(2.3))
c1.paragraphs[0].add_run().add_picture(r"C:\Users\tishy\Documents\Honors\docs\example_eczema.jpg", width=Inches(2.3))
cap0, cap1 = ex_table.rows[1].cells
cap0.paragraphs[0].add_run("Normal/ example — stock photo, studio background").italic = True
cap1.paragraphs[0].add_run("Eczema/ example — clinical macro photo, DermNet watermark visible").italic = True

p(
    "This is a classic shortcut-learning / confounding-variable risk: a classifier can reach very "
    "high accuracy by learning 'stock-photo composition vs. clinical macro-photo (with watermark)' "
    "rather than 'eczema vs. healthy skin.' This would explain why accuracy is high enough to be "
    "suspicious (95.66% test accuracy on a first, lightly-tuned baseline). It does not mean the "
    "model has learned nothing about eczema, but the current number should not be reported or relied "
    "on until this is ruled out."
)
p("Recommended before trusting or reporting this result further:", bold=True)
bullet("Re-crop both classes to skin-region-only patches with consistent framing, removing "
       "background/composition as a usable cue.")
bullet("Remove or inpaint watermarked images (search for the DermNet mark specifically, and audit "
       "for other source watermarks).")
bullet("Re-run evaluation on a small manually-curated 'clean-style' held-out set (same photography "
       "style for both classes) to see whether accuracy holds up — a large drop would confirm shortcut "
       "learning.")
bullet("Consider a same-style source for the Normal class (e.g. normal-skin crops from the same "
       "clinical atlases used for Eczema, rather than stock photography) for a fairer comparison.")

# ===================== SECTION 4: STAGE A DATASET =====================
h1("4. Stage A — Motion Dataset (WISDM)")

h2("4.1 Dataset choice")
p(
    "The project report's own next-steps list named CAPTURE-24 as the priority motion dataset. "
    "In this session, WISDM was chosen instead as a faster first dataset to get a pipeline running: "
    "smaller download, simple per-line CSV format, and labelled everyday activities from wrist and "
    "phone accelerometer/gyroscope sensors (51 subjects, 18 activities). WESAD was considered and "
    "rejected as the primary choice for Stage A specifically, because its labels are stress/affect "
    "conditions (baseline, stress, amusement, meditation) rather than activity/scratch labels, making "
    "it a weaker negative-class source for a scratch-vs-non-scratch classifier than WISDM's diverse "
    "labelled activities. CAPTURE-24 remains the larger, more free-living dataset to pursue next, "
    "per the original report."
)

h2("4.2 Download")
p(
    "Source: UCI Machine Learning Repository (WISDM Smartphone and Smartwatch Activity and "
    "Biometrics Dataset). Downloaded to dataset/WISDM/, extracted (~897MB), and the intermediate "
    ".zip files were deleted after extraction to save disk space."
)
bullet("51 subjects x 2 devices (phone, watch) x 2 sensors (accelerometer, gyroscope) = 204 raw files")
bullet("Raw format per line: subject_id, activity_code, timestamp, x, y, z;")
bullet("18 labelled activities (walking, jogging, stairs, sitting, standing, typing, teeth-brushing, "
       "eating variants, drinking, kicking, catch, dribbling, writing, clapping, folding)")
p(
    "Note: 'teeth' (brushing teeth) is the closest proxy in this dataset to a repetitive hand-motion "
    "gesture, worth keeping in mind when this data is used to prototype scratch-vs-non-scratch "
    "features later — none of the activities are scratch itself."
)

h2("4.3 Cleaning process")
p(
    "A cleaning script (scripts/clean_wisdm.py) was written and run over all 204 raw files. For "
    "every row it:"
)
bullet("Validated the line has exactly 6 fields (subject, activity, timestamp, x, y, z).")
bullet("Validated the subject id in the row matches the subject id in the filename.")
bullet("Validated the activity code exists in activity_key.txt.")
bullet("Validated timestamp parses as an integer and x/y/z parse as finite floats (catches NaN/inf).")
bullet("Dropped duplicate rows sharing an identical timestamp within a file.")
bullet("Flagged (but kept) readings with |value| > 50 in any axis, as a statistical outlier report "
       "rather than an assumed error — large values can reflect genuine sudden motion.")
p("Cleaned files were written to a parallel directory (dataset/WISDM/raw_clean/) rather than overwriting the originals.")

h2("4.4 Results")
table_from_rows(
    ["Metric", "Value"],
    [
        ["Files processed", "204"],
        ["Rows kept", "15,364,411"],
        ["Rows dropped (duplicate timestamp)", "266,015"],
        ["Malformed / wrong field count", "0"],
        ["Invalid activity codes", "0"],
        ["Subject-id mismatches", "0"],
        ["Non-finite values", "0"],
        ["Outlier readings flagged (kept)", "1,656"],
    ],
)
p(
    "The dataset was largely clean already — the only systematic issue found was duplicate-timestamp "
    "rows (about 1.7% of all rows), which is a known artifact of this dataset's logging. No "
    "corruption, malformed lines, or mismatched subjects were found."
)
p("Breakdown by device/sensor:")
table_from_rows(
    ["Device", "Sensor", "Files", "Rows Kept", "Rows Dropped", "Outliers Flagged"],
    [
        ["Phone", "Accelerometer", 51, "4,734,730", "69,673", "1,146"],
        ["Phone", "Gyroscope", 51, "3,544,385", "64,250", "0"],
        ["Watch", "Accelerometer", 51, "3,710,982", "66,064", "510"],
        ["Watch", "Gyroscope", 51, "3,374,314", "66,028", "0"],
    ],
)
p(
    "Outputs: dataset/WISDM/raw_clean/<device>/<sensor>/ (cleaned per-subject files, same format as "
    "source), dataset/WISDM/clean_report.csv (per-file kept/dropped counts by reason), "
    "dataset/WISDM/clean_report_lines.csv (example dropped lines, capped at 20 per file)."
)

# ===================== SECTION 5: NEXT STEPS =====================
h1("5. Current Status and Next Steps")
h2("5.1 Completed this session")
bullet("Cleaned and split the Stage B image dataset (Eczema/Normal).")
bullet("Trained and evaluated a first Stage B baseline model (ResNet18 transfer learning, 95.66% test accuracy).")
bullet("Downloaded and cleaned the WISDM motion dataset for Stage A.")

h2("5.2 Not yet done")
bullet("HIGH PRIORITY: investigate and fix the likely shortcut-learning issue in the Stage B image "
       "dataset (Section 3.4) before trusting or reporting the 95.66% test accuracy figure.")
bullet("Stage A model itself has not been trained yet — WISDM is cleaned and ready, but no "
       "scratch/motion classifier has been built.")
bullet("CAPTURE-24 (the report's originally named priority dataset) has not been downloaded.")
bullet("Stage C fusion model has not been started.")
bullet("Ethics/IRB paperwork status is unchanged from the original report and remains the item on "
       "the critical path for real patient data.")
bullet("Stage B model above is a first baseline (frozen backbone, 10 epochs); further iteration "
       "(unfreezing, longer training) was not yet attempted.")
p(
    "See Section 7 for an updated, literature-informed recommendation on how these remaining items "
    "should be sequenced and framed, following review of Maulana et al. (2024) in Section 6, and "
    "Section 8 for a follow-up experiment that substantially strengthens the shortcut-learning "
    "concern above.",
    italic=True,
)

# ===================== SECTION 6: NEW LITERATURE =====================
h1("6. New Literature Reviewed — Maulana et al. (2024)")
p(
    "Maulana, A. et al., \"Enhanced Prediction of Atopic Dermatitis Severity Using Advanced Machine "
    "Learning Techniques,\" 2024 International Conference on Electrical Engineering and Informatics "
    "(ICELTICs). (PDF supplied by Lilac; saved at project root.)",
    italic=True,
)

h2("6.1 What the paper did")
bullet("Dataset: 3,037 clinical close-up photos of AD lesions from 250 patients at a hospital in "
       "Banda Aceh, Indonesia, all captured with the same 12-megapixel smartphone protocol under "
       "dermatologist supervision, then manually cropped to the lesion region.")
bullet("Labels: each image was scored using the objective SCORAD index by a dermatologist and binned "
       "into 4 severity classes: None, Mild, Moderate, Severe (not a binary eczema/normal split).")
bullet("Features: no deep learning / CNN embeddings at all. Instead, 90 hand-engineered colour "
       "features (mean + standard deviation) extracted per image across many colour spaces: RGB, "
       "normalized RGB, YCbCr, HSV, HLS, CIE XYZ/LAB/LUV/LCH, Opponent, CMY, YUV, YIQ, YDbDr, YPbPr.")
bullet("Models: these 90 features were fed into four classical gradient-boosting / ensemble models "
       "— XGBoost, CatBoost, LightGBM, and Random Forest — trained to predict the 4-class severity "
       "label.")
bullet("Results: all four models scored 93-95% accuracy; LightGBM was best (95.07% accuracy, 95.10% "
       "precision, 95.07% recall, 95.07% F1), with XGBoost close behind (94.74%). SHAP analysis showed "
       "standard-deviation-of-hue/chrominance features (e.g. STD_CIE_LCH_H, STD_YIQ_Q, STD_HLS_L) "
       "were the most predictive, i.e. colour variability within a lesion, not just its average colour.")
bullet("Confusion matrix: nearly all errors were between the adjacent Mild and Moderate classes; "
       "None and Severe were classified very cleanly. The authors attribute this to genuinely "
       "overlapping clinical presentation and SCORAD-labelling subjectivity between those two grades.")
bullet("The authors' own stated limitations: single-site population (Aceh only, no external "
       "validation), no explicit dataset-leakage/shortcut check, potential overfitting in the "
       "boosting models, and ethical/privacy considerations for clinical deployment.")

h2("6.2 Why this is relevant to this project")
p(
    "This paper is a close match to Stage B of this project (image-based lesion severity scoring) "
    "and is useful in three concrete ways:",
)
bullet(
    "It validates the achievable accuracy range for a properly controlled image-severity model: "
    "93-95%, on a dataset where every image (across all four severity classes) came from the same "
    "camera, same protocol, same clinical setting. There is no possible shortcut based on photo "
    "source or style, because the source is identical across classes. This is the key structural "
    "difference from this project's current Stage B dataset (Section 3.4), where the two classes "
    "come from different sources (stock photography vs. DermNet clinical archive) and the "
    "suspiciously similar 95-97% figure is much more likely to be picking up that difference. In "
    "other words: this paper is evidence for what the number should look like when it is genuine, "
    "and a concrete argument for why this project's dataset needs the same-source discipline before "
    "its own accuracy figure can be trusted."
)
bullet(
    "It suggests the project's Stage B target should arguably be graded severity (None/Mild/"
    "Moderate/Severe, matching SCORAD) rather than a binary Eczema-vs-Normal decision, since severity "
    "grading is what feeds meaningfully into a fusion model that is also trying to track how bad a "
    "flare is over time, not just whether one is present."
)
bullet(
    "It offers a second, complementary modelling approach to the ResNet18 transfer-learning "
    "pipeline already built: hand-engineered, interpretable colour features plus gradient boosting. "
    "This is worth having as a comparison/ensemble arm, and its use of SHAP gives a template for "
    "explaining Stage B's predictions, which strengthens the eventual paper regardless of which "
    "model is used in production."
)

# ===================== SECTION 7: RECOMMENDED PATH FORWARD =====================
h1("7. Recommended Path Forward")

h2("7.1 Immediate priority (carried over): fix the Stage B dataset")
p(
    "Unchanged from Section 3.4 and still blocking: re-source or re-curate the Eczema/Normal images "
    "so both classes share a consistent capture style (ideally by adopting the same-source "
    "discipline Maulana et al. used), remove watermarked images, and re-evaluate before trusting any "
    "accuracy number from this dataset."
)

h2("7.2 Reframe Stage B as graded severity, not binary")
p(
    "Move from Eczema-vs-Normal toward a SCORAD-informed severity scale (even a simplified 3-class "
    "Mild/Moderate/Severe, if a 4th 'None' class is hard to source cleanly). This aligns Stage B's "
    "output with what Stage C (fusion) actually needs: a severity signal to combine with motion data, "
    "not just a presence/absence flag."
)

h2("7.3 Add a second Stage B model: engineered colour features + gradient boosting")
p(
    "Alongside the existing ResNet18 transfer-learning model, implement the colour-feature-extraction "
    "+ LightGBM/XGBoost pipeline described in Maulana et al. (90 features across the colour spaces "
    "listed in Section 6.1). This is cheap to build (classical ML, CPU-friendly, no GPU needed), gives "
    "a second opinion to sanity-check the CNN model against, and produces SHAP feature-importance "
    "explanations that are valuable for the paper's discussion section."
)

h2("7.4 Build the Stage A scratch/motion model")
p(
    "The WISDM dataset is cleaned and ready (Section 4). Next step is a first classifier "
    "distinguishing scratch-like repetitive hand motion from other activities using the accelerometer/"
    "gyroscope windows, likely starting from simple windowed statistical features (mean, std, "
    "zero-crossing rate, dominant frequency) feeding a classical classifier (Random Forest / "
    "gradient boosting, matching the interpretable approach used successfully in Section 6), before "
    "considering a deep sequence model."
)

h2("7.5 Build Stage C: the fusion model")
p(
    "Once Stage A and a trustworthy Stage B exist, combine per-night/per-session outputs (scratch "
    "frequency/intensity from Stage A, lesion severity score from Stage B) into a single fusion model. "
    "Gradient boosting (XGBoost/LightGBM) is a reasonable first choice here too, given its strong, "
    "interpretable performance in both this project's own use case and in the literature reviewed."
)

h2("7.6 Suggested framing for the honors paper")
p(
    "Given the literature landscape, a strong, defensible framing for the paper is:", bold=True
)
p(
    "\"A multi-sensor wearable system for at-home atopic dermatitis monitoring, fusing wrist-worn "
    "motion-based scratch detection with smartphone image-based severity scoring.\""
)
p("This framing is distinct from existing published work in two specific ways worth stating explicitly in the paper:")
bullet(
    "Existing image-severity work (including Maulana et al. 2024, and the CNN literature it cites) "
    "is single-modality and clinic-captured: a dermatologist or hospital protocol takes the photo. "
    "This project's contribution is combining that with a continuous, passive, at-home behavioural "
    "signal (scratching, captured via wearable motion sensing) that clinic photography cannot see at "
    "all — itch/scratch behaviour between clinic visits."
)
bullet(
    "None of the datasets used so far (WISDM for motion, the Eczema/Normal image set) were collected "
    "with this fused use case in mind, which should be stated as a limitation up front, the same way "
    "Maulana et al. transparently flagged their own single-site, single-population limitation. The "
    "paper's contribution at this stage is the pipeline and fusion methodology, validated on proxy "
    "datasets, with a clear path to real multi-modal patient data pending ethics approval."
)
bullet(
    "Reporting the shortcut-learning finding from Section 3.4 itself, and how it was diagnosed and "
    "addressed, is legitimate methodological content for the paper — it demonstrates the kind of "
    "dataset-validity rigor that Maulana et al. explicitly identify as missing from their own and "
    "prior work (no external validation, no explicit leakage check)."
)

# ===================== SECTION 8: COLOUR-FEATURE + LIGHTGBM EXPERIMENT =====================
h1("8. Second Stage B Model — Colour Features + LightGBM (Testing the Shortcut-Learning Hypothesis)")
p(
    "Following the recommendation in Section 7.3, the colour-feature-extraction + gradient-boosting "
    "pipeline from Maulana et al. (2024) was implemented and run on this project's own Eczema/Normal "
    "image set, using the exact same train/val/test split as the ResNet18 model (Section 3), so the "
    "two approaches are directly comparable on identical held-out images."
)

h2("8.1 Method")
bullet(
    "scripts/extract_color_features.py computes the same 90 features as the paper (mean + standard "
    "deviation of each channel, across 15 colour representations: RGB, normalized RGB, YCbCr, HSV, "
    "HLS, CIE XYZ, CIE LAB, CIE LCH, CIE LUV, Opponent, CMY, YUV, YIQ, YDbDr, YPbPr). Images are "
    "resized to 128x128 before extraction; every image in manifest_clean.csv (2,757 images) was "
    "processed successfully."
)
bullet(
    "scripts/train_stage_b_lightgbm.py trains a LightGBM binary classifier (Eczema vs. Normal) on "
    "these 90 features using manifest_train.csv, with manifest_val.csv for early stopping, and "
    "reports final metrics on manifest_test.csv — the same 415-image test set used for the CNN."
)
p(
    "Note: this reproduces the paper's model architecture and feature-engineering approach, but not "
    "its label set — the paper predicted 4-class SCORAD severity from clinician labels, which this "
    "project's dataset does not have. Here the same features and model type were applied to this "
    "project's actual available label (Eczema vs. Normal), which is a fair reuse of the modelling "
    "approach even though the task is not identical."
)

h2("8.2 Results")
table_from_rows(
    ["Metric", "ResNet18 (CNN)", "Colour Features + LightGBM"],
    [
        ["Test set size", "415", "415"],
        ["Test accuracy", "95.66%", "95.18%"],
        ["Precision (Eczema)", "92.64%", "93.72%"],
        ["Recall (Eczema)", "99.53%", "97.21%"],
        ["F1 (Eczema)", "95.96%", "95.43%"],
    ],
)
p("LightGBM confusion matrix (rows = true label, columns = predicted label):")
table_from_rows(
    ["", "Predicted Normal", "Predicted Eczema"],
    [
        ["True Normal", "186 (TN)", "14 (FP)"],
        ["True Eczema", "6 (FN)", "209 (TP)"],
    ],
)
doc.add_picture(r"C:\Users\tishy\Documents\Honors\docs\confusion_matrix_lgbm.png", width=Inches(4.0))

h2("8.3 Top contributing features")
table_from_rows(
    ["Rank", "Feature", "Gain"],
    [
        [1, "MEAN_XYZ_2 (mean of CIE-XYZ Z channel — overall blue/luminance level)", "13,608.8"],
        [2, "STD_XYZ_2", "1,637.1"],
        [3, "MEAN_RGB_1 (mean Green channel)", "1,575.9"],
        [4, "MEAN_XYZ_1", "1,307.8"],
        [5, "MEAN_CMY_1", "626.4"],
    ],
)

h2("8.4 Interpretation — this strongly confirms the shortcut-learning concern")
para = p(
    "A model using only whole-image colour averages and variances — no shape, no texture, no spatial "
    "structure, no lesion boundary information at all — matches the CNN's test accuracy almost "
    "exactly (95.18% vs. 95.66%). This is strong quantitative evidence for the shortcut-learning "
    "hypothesis raised in Section 3.4.",
    bold=True,
)
p(
    "The single most important feature by a wide margin is MEAN_XYZ_2 (mean overall brightness/blue "
    "level of the whole image), followed by MEAN_RGB_1 (mean green channel) — these are global "
    "exposure/colour-balance statistics, exactly the kind of signal that would differ systematically "
    "between Shutterstock-style studio photography (Normal/) and DermNet-style clinical macro "
    "photography (Eczema/) regardless of skin condition. A model that needs to look at actual lesion "
    "texture or shape would be expected to rely on standard-deviation features tied to local colour "
    "variation within the lesion (as in Maulana et al.'s SHAP results, Section 6.1) rather than the "
    "image's overall mean colour."
)
p(
    "Combined with the visual/filename-pattern evidence already documented in Section 3.4, this "
    "experiment upgrades the shortcut-learning concern from 'likely' to 'well-supported': two "
    "completely different model families (a pretrained CNN and a gradient-boosted tree ensemble on "
    "hand-engineered colour statistics) achieve near-identical, suspiciously high accuracy using "
    "features that plausibly capture photo source/lighting rather than eczema presence. The dataset "
    "fix recommended in Section 3.4 and Section 7.1 (same-source imagery for both classes) should be "
    "treated as a prerequisite, not an optional improvement, before either model's accuracy is "
    "reported as a genuine eczema-detection result."
)

doc.save(OUT)
print(f"Saved: {OUT}")
